
# =========================================================
# VidIntel — YouTube + RSS Article Discovery (Multilingual)
# Streamlit Application (Refactored with Looser Article Search)
# =========================================================

import re
import unicodedata
import requests
import pandas as pd
import streamlit as st
from datetime import datetime, timedelta, timezone
from bs4 import BeautifulSoup
from urllib.parse import urljoin

# Optional deps
try:
    from langdetect import detect
except Exception:
    detect = None

try:
    from docx import Document
except Exception:
    Document = None


# -------------------------
# Streamlit Configuration
# -------------------------
st.set_page_config(
    page_title="VidIntel",
    page_icon="📺",
    layout="wide"
)

API_KEY = st.secrets.get("YOUTUBE_API_KEY", None)

if not API_KEY:
    st.error(
        "❌ YouTube API key not found.\n\n"
        "Add it in **Streamlit → App settings → Secrets** as:\n"
        "`YOUTUBE_API_KEY = \"your_key_here\"`"
    )
    st.stop()

SEARCH_URL = "https://www.googleapis.com/youtube/v3/search"
VIDEOS_URL = "https://www.googleapis.com/youtube/v3/videos"


# =========================================================
# Human-friendly labels for Languages & Regions
# =========================================================

LANG_LABELS = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "ru": "Russian",
    "ar": "Arabic",
    "fa": "Persian (Farsi)",
    "ja": "Japanese",
    "ko": "Korean",
}

# Options for selectboxes: list of (Label, Code) tuples
VIDEO_LANG_OPTIONS = [(name, code) for code, name in LANG_LABELS.items()]
ARTICLE_LANG_OPTIONS = [(name, code) for code, name in LANG_LABELS.items()]

REGIONS = [
    ("Any region", None),

    # English-dominant
    ("US (United States)", "US"),
    ("GB (United Kingdom)", "GB"),
    ("CA (Canada)", "CA"),
    ("AU (Australia)", "AU"),

    # Spanish-speaking
    ("ES (Spain — Europe)", "ES"),
    ("MX (Mexico — North America)", "MX"),
    ("AR (Argentina — South America)", "AR"),
    ("CO (Colombia — South America)", "CO"),
    ("CL (Chile — South America)", "CL"),
    ("PE (Peru — South America)", "PE"),
    ("DO (Dominican Republic — Caribbean)", "DO"),

    # Russian-speaking
    ("RU (Russia)", "RU"),
    ("KZ (Kazakhstan)", "KZ"),
    ("UA (Ukraine)", "UA"),

    # Arabic-speaking
    ("EG (Egypt)", "EG"),
    ("SA (Saudi Arabia)", "SA"),
    ("AE (United Arab Emirates)", "AE"),

    # East Asia
    ("JP (Japan)", "JP"),
    ("KR (South Korea)", "KR"),
]


# =========================================================
# Keyword Synonyms (language-aware) for Articles
# Expand as needed for your common topics.
# Keys should be lowercase.
# =========================================================

ARTICLE_KEYWORD_SYNONYMS = {
    "en": {
        "economy": ["economy", "economic", "economics", "finance", "inflation", "market", "markets", "gdp"],
        "article": ["article", "report", "coverage", "story", "news"],
        "ai": ["ai", "artificial intelligence", "machine learning", "ml"],
    },
    "es": {
        "economy": ["economía", "economico", "económico", "economicos", "económicos",
                    "finanzas", "inflación", "mercado", "mercados", "pib"],
        "article": ["artículo", "articulo", "cobertura", "noticia", "informe", "reporte"],
        "ai": ["ia", "inteligencia artificial", "aprendizaje automático", "aprendizaje estadístico"],
    },
    "fr": {
        "economy": ["économie", "economie", "économique", "economique", "finance", "inflation", "marché", "marchés"],
    },
    # Add more languages/topics as needed…
}


# =========================================================
# Utility: Normalization for accent-insensitive match
# =========================================================
def normalize_text(text: str) -> str:
    """
    Lowercase, strip accents/diacritics, and collapse whitespace.
    """
    if not text:
        return ""
    # Normalize to NFKD and remove combining marks (accents)
    nfkd = unicodedata.normalize("NFKD", text)
    base = "".join(ch for ch in nfkd if not unicodedata.combining(ch))
    base = base.lower()
    return " ".join(base.split())


# =========================================================
# YOUTUBE SEARCH (existing logic, with small guards)
# =========================================================

def iso8601_to_seconds(iso):
    m = re.match(r"^PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?$", iso or "")
    if not m:
        return 0
    h = int(m.group(1) or 0)
    mnt = int(m.group(2) or 0)
    s = int(m.group(3) or 0)
    return h * 3600 + mnt * 60 + s


def compute_published_after(choice):
    now = datetime.now(timezone.utc)
    if choice == "No filter":
        return None
    if choice == "Last 24 hours":
        return now - timedelta(hours=24)
    if choice == "Last 72 hours":
        return now - timedelta(hours=72)
    if choice == "Last 7 days":
        return now - timedelta(days=7)
    if choice == "This month":
        return now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    if choice == "This year":
        return now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
    return None


def search_youtube(query, lang, region, max_results, published_after):
    params = {
        "part": "snippet",
        "q": query,
        "type": "video",
        "maxResults": max_results,
        "relevanceLanguage": lang,
        "key": API_KEY,
    }
    if region:
        params["regionCode"] = region
    if published_after:
        params["publishedAfter"] = published_after.isoformat().replace("+00:00", "Z")
        params["order"] = "date"

    r = requests.get(SEARCH_URL, params=params, timeout=30)
    r.raise_for_status()
    return r.json().get("items", [])


def fetch_video_details(video_ids):
    if not video_ids:
        return {}
    params = {
        "part": "contentDetails,statistics",
        "id": ",".join(video_ids),
        "key": API_KEY,
    }
    r = requests.get(VIDEOS_URL, params=params, timeout=30)
    r.raise_for_status()

    details = {}
    for item in r.json().get("items", []):
        vid = item.get("id")
        dur_iso = item.get("contentDetails", {}).get("duration")
        duration_sec = iso8601_to_seconds(dur_iso) if dur_iso else 0
        details[vid] = {
            "duration_sec": duration_sec,
            "views": int(item.get("statistics", {}).get("viewCount", 0)),
        }
    return details


def strict_language_filter(items, target_lang):
    if not detect:
        return items
    filtered = []
    for it in items:
        text = it["snippet"]["title"] + " " + it["snippet"].get("description", "")
        try:
            if detect(text) == target_lang:
                filtered.append(it)
        except Exception:
            pass
    return filtered


def youtube_to_dataframe(items, details):
    rows = []
    for it in items:
        s = it["snippet"]
        vid = it["id"]["videoId"]
        mins = round(details.get(vid, {}).get("duration_sec", 0) / 60, 1)
        rows.append({
            "Title": s["title"],
            "Channel": s["channelTitle"],
            "Published": s["publishedAt"][:10],
            "Duration (min)": mins,
            "URL": f"https://www.youtube.com/watch?v={vid}",
        })
    return pd.DataFrame(rows)


def export_videos_txt(df):
    lines = []
    for i, r in df.iterrows():
        lines += [
            f"{i+1}. {r['Title']}",
            f"   Channel: {r['Channel']}",
            f"   Published: {r['Published']}",
            f"   Duration: {r['Duration (min)']} min",
            f"   URL: {r['URL']}",
            ""
        ]
    return "\n".join(lines)


def export_videos_docx(df):
    if not Document:
        return None
    doc = Document()
    doc.add_heading("VidIntel YouTube Results", level=1)
    for _, r in df.iterrows():
        p = doc.add_paragraph()
        p.add_run(r["Title"] + "\n").bold = True
        p.add_run(r["Channel"] + "\n")
        p.add_run(f"Published: {r['Published']} | {r['Duration (min)']} min\n")
        p.add_run(r["URL"])
    return doc


# =========================================================
# ARTICLE SEARCH (RSS) — Looser, smarter matching
# =========================================================

CURATED_RSS = {
    "en": {
        "BBC": "https://feeds.bbci.co.uk/news/rss.xml",
        "Reuters": "https://feeds.reuters.com/reuters/worldNews",
        "AP News": "https://apnews.com/rss",
    },
    "es": {
        "El País": "https://feeds.elpais.com/mrss-s/pages/ep/site/elpais.com/portada",
        "La Nación (AR)": "https://www.lanacion.com.ar/arc/outboundfeeds/rss/",
        "El Universal (MX)": "https://www.eluniversal.com.mx/rss.xml",
    },
    "fr": {
        "Le Monde": "https://www.lemonde.fr/rss/une.xml",
        "France 24": "https://www.france24.com/fr/rss",
    },
    "de": {
        "Der Spiegel": "https://www.spiegel.de/international/index.rss",
        "DW": "https://rss.dw.com/xml/rss-en-all",
    },
    "ru": {
        "Meduza": "https://meduza.io/rss/all",
        "Kommersant": "https://www.kommersant.ru/RSS/news.xml",
    },
    "ar": {
        "Al Jazeera": "https://www.aljazeera.net/aljazeera/rss",
        "Al Arabiya": "https://english.alarabiya.net/feed/rss",
    },
    "fa": {"BBC Persian": "https://feeds.bbci.co.uk/persian/rss.xml"},
    "ja": {"NHK": "https://www3.nhk.or.jp/rss/news/cat0.xml"},
    "ko": {"Yonhap": "https://en.yna.co.kr/rss/all.xml"},
}


def discover_rss_feeds(site_url, timeout=10):
    feeds = []
    try:
        r = requests.get(site_url, timeout=timeout, headers={"User-Agent": "VidIntelBot/1.0"})
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        for link in soup.find_all("link", type=["application/rss+xml", "application/atom+xml"]):
            href = link.get("href")
            if href:
                feeds.append(urljoin(site_url, href))
    except Exception:
        pass

    # fallback guesses
    for path in ["/rss", "/rss.xml", "/feed"]:
        feeds.append(urljoin(site_url, path))

    return list(set(feeds))


def get_article_sources(language_code, user_url=None):
    sources = {}
    sources.update(CURATED_RSS.get(language_code, {}))

    if user_url:
        attempts = discover_rss_feeds(user_url)
        for i, feed in enumerate(attempts[:3]):
            sources[f"User Source {i+1}"] = feed

    return sources


@st.cache_data(ttl=600, show_spinner=False)
def fetch_articles(
    sources: dict,
    days_back: int = 7,
    keyword: str | None = None,
    article_lang: str = "en",
    loose_filter: bool = False,
):
    """
    Fetch and filter articles from the given RSS/Atom sources.
    - Full-text matching: title + summary + <content:encoded> when available
    - Accent-insensitive and case-insensitive
    - Language-aware synonyms expansion
    - Optional 'loose_filter' to ignore keyword matching entirely
    """
    from dateutil import parser as dateparser

    rows = []
    cutoff_ts = datetime.now(timezone.utc).timestamp() - days_back * 86400

    # Prepare keyword set:
    # Split user keyword on commas or whitespace; expand with language-aware synonyms
    expanded_terms = set()
    if keyword:
        # Split on commas first; then split each on whitespace
        raw_parts = []
        for chunk in keyword.split(","):
            raw_parts.extend(chunk.strip().split())

        for term in raw_parts:
            t_norm = normalize_text(term)
            if not t_norm:
                continue

            # Add the normalized user term
            expanded_terms.add(t_norm)

            # Add synonyms, if any
            lang_map = ARTICLE_KEYWORD_SYNONYMS.get(article_lang, {})
            if t_norm in lang_map:
                for syn in lang_map[t_norm]:
                    expanded_terms.add(normalize_text(syn))

    # Pull feeds
    for name, url in sources.items():
        try:
            resp = requests.get(url, timeout=12, headers={"User-Agent": "VidIntelBot/1.0"})
            resp.raise_for_status()
            xml_text = resp.text
        except Exception:
            continue

        soup_xml = BeautifulSoup(xml_text, "xml")

        # Prefer RSS <item>; fall back to Atom <entry>
        entries = soup_xml.find_all("item")
        is_atom = False
        if not entries:
            entries = soup_xml.find_all("entry")
            is_atom = True

        for entry in entries:
            if not is_atom:
                title = (entry.title.text if entry.title else "").strip()
                link = (entry.link.text if entry.link else "").strip()
                pub = (entry.pubDate.text if entry.pubDate else "").strip()
                summary_html = (entry.description.text if entry.description else "")
                # Full text (if present via content:encoded)
                content_node = entry.find("content:encoded")
                full_text = content_node.text if content_node else ""
            else:
                title = (entry.title.text if entry.title else "").strip()
                # Atom link: pick rel="alternate" if present; else first href
                link_tag = None
                for l in entry.find_all("link"):
                    if l.get("rel") == "alternate" and l.get("href"):
                        link_tag = l
                        break
                if not link_tag:
                    link_tag = entry.find("link")
                link = (link_tag.get("href") if link_tag and link_tag.has_attr("href") else "").strip()
                pub = (entry.updated.text if entry.find("updated") else
                       entry.published.text if entry.find("published") else "").strip()
                summary_html = (entry.summary.text if entry.find("summary") else "")
                content_node = entry.find("content")
                full_text = content_node.text if content_node else ""

            # Parse time and apply cutoff
            try:
                ts = dateparser.parse(pub).timestamp() if pub else None
            except Exception:
                ts = None
            if ts and ts < cutoff_ts:
                continue

            # Build full blob and normalize (accent-insensitive)
            summary_txt = BeautifulSoup(summary_html, "html.parser").get_text().strip()
            blob_norm = normalize_text(f"{title} {summary_txt} {full_text}")

            # Apply keyword logic
            if keyword and not loose_filter and expanded_terms:
                if not any(term in blob_norm for term in expanded_terms):
                    continue
            # If loose_filter is True OR no keyword provided, accept based on date alone

            rows.append({
                "Title": title,
                "Source": name,
                "Published": pub,
                "URL": link,
                "Summary": summary_txt
            })

    # Sort newest first (try parse dates; push invalid to end)
    def sort_key(row):
        try:
            from dateutil import parser as dateparser  # local import for cache hashing
            return dateparser.parse(row["Published"])
        except Exception:
            return datetime.min.replace(tzinfo=timezone.utc)

    rows.sort(key=sort_key, reverse=True)
    return pd.DataFrame(rows)


def export_articles_txt(df):
    parts = []
    for _, r in df.iterrows():
        parts.append(
            f"{r['Title']} ({r['Source']})\n{r['Published']}\n{r['URL']}\n\n{r['Summary']}\n"
        )
    return "\n".join(parts)


def export_articles_docx(df):
    if not Document:
        return None
    doc = Document()
    doc.add_heading("VidIntel Article Results", level=1)
    for _, r in df.iterrows():
        p = doc.add_paragraph()
        p.add_run(r["Title"] + "\n").bold = True
        p.add_run(f"{r['Source']} — {r['Published']}\n")
        p.add_run(r["URL"] + "\n")
        if r.get("Summary"):
            p.add_run("\nSummary:\n")
            p.add_run(r["Summary"] + "\n")
    return doc


# =========================================================
# UI — Main App
# =========================================================

st.title("🎯 VidIntel — YouTube + Article Discovery")

with st.sidebar:
    st.header("🔎 Search Filters")

    content_type = st.selectbox(
        "Content Type",
        ["Videos", "Articles", "Both"]
    )

    keyword = st.text_input("Keyword / Query", "", placeholder="e.g., economy, AI, defense")

    # -------------------------
    # 🎬 Video Filters
    # -------------------------
    video_lang_label, lang = st.selectbox(
        "Video Language",
        VIDEO_LANG_OPTIONS,
        index=0,
        format_func=lambda x: x[0]
    )

    video_region_label, region = st.selectbox(
        "Region",
        REGIONS,
        index=0,
        format_func=lambda x: x[0]
    )

    date_filter = st.selectbox(
        "Video Date Range",
        ["No filter", "Last 24 hours", "Last 72 hours", "Last 7 days", "This month", "This year"]
    )

    duration_filter = st.selectbox(
        "Duration",
        ["Any", "< 4 min", "5–10 min", "> 10 min"]
    )

    strict_lang = st.checkbox("Strict language filtering")
    max_results = st.slider("Max video results", 5, 50, 20, step=5)

    st.markdown("---")

    # -------------------------
    # 📰 Article Filters
    # -------------------------
    article_lang_label, article_lang = st.selectbox(
        "Article Language",
        ARTICLE_LANG_OPTIONS,
        index=0,
        format_func=lambda x: x[0]
    )

    days_back = st.slider("Article Days Back", 1, 30, 7)
    loose_filter = st.checkbox("Loosen article filtering (ignore keyword match)", value=False)
    user_site = st.text_input("Auto-discover RSS from site (optional)", "", placeholder="https://example.com")

    run = st.button("🔍 Search", type="primary")


# =========================================================
# Execution
# =========================================================

videos_df = None
articles_df = None
article_sources_count = 0

if run:
    # -------------------------
    # VIDEOS
    # -------------------------
    if content_type in ["Videos", "Both"]:
        if not keyword.strip():
            st.warning("Enter a keyword to search YouTube videos (Videos tab).")
        else:
            with st.spinner("Searching YouTube…"):
                published_after = compute_published_after(date_filter)
                items = search_youtube(keyword, lang, region, max_results, published_after)

                ids = [it["id"]["videoId"] for it in items]
                details = fetch_video_details(ids)

                # Duration filtering
                if duration_filter != "Any":
                    def ok(d):
                        m = d / 60
                        return (
                            (duration_filter == "< 4 min" and m < 4) or
                            (duration_filter == "5–10 min" and 5 <= m <= 10) or
                            (duration_filter == "> 10 min" and m > 10)
                        )
                    items = [
                        it for it in items
                        if ok(details.get(it["id"]["videoId"], {}).get("duration_sec", 0))
                    ]

                if strict_lang:
                    items = strict_language_filter(items, lang)

                videos_df = youtube_to_dataframe(items, details)

    # -------------------------
    # ARTICLES
    # -------------------------
    if content_type in ["Articles", "Both"]:
        with st.spinner("Fetching articles…"):
            sources = get_article_sources(article_lang, user_site or None)
            article_sources_count = len(sources)

            articles_df = fetch_articles(
                sources=sources,
                days_back=days_back,
                keyword=keyword if keyword.strip() else None,
                article_lang=article_lang,
                loose_filter=loose_filter
            )


# =========================================================
# OUTPUT TABS
# =========================================================

tabs = st.tabs(["📺 Videos", "📰 Articles"])

# -------------------------
# VIDEOS TAB
# -------------------------
with tabs[0]:
    st.markdown(
        f"##### Filters: **{video_lang_label}** | **{video_region_label or 'Any region'}** | **{date_filter}** | **{duration_filter}**"
    )

    if videos_df is None:
        st.info("⚠️ Choose 'Videos' or 'Both' to search for YouTube content.")
    elif videos_df.empty:
        st.warning("No YouTube results.")
    else:
        st.success(f"Found {len(videos_df)} videos")

        st.dataframe(
            videos_df,
            use_container_width=True,
            column_config={
                "URL": st.column_config.LinkColumn("YouTube Link")
            }
        )

        from io import BytesIO
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "⬇️ CSV",
                videos_df.to_csv(index=False).encode("utf-8"),
                "videos.csv",
                "text/csv"
            )
        with col2:
            st.download_button(
                "⬇️ TXT",
                export_videos_txt(videos_df).encode("utf-8"),
                "videos.txt",
                "text/plain"
            )
        with col3:
            doc = export_videos_docx(videos_df)
            if doc:
                buf = BytesIO()
                doc.save(buf)
                st.download_button(
                    "⬇️ Word",
                    buf.getvalue(),
                    "videos.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# -------------------------
# ARTICLES TAB
# -------------------------
with tabs[1]:
    st.markdown(
        f"##### Filters: **{article_lang_label}** | Days back: **{days_back}** | Sources: **{article_sources_count}** | Loosen: **{'On' if loose_filter else 'Off'}**"
    )

    if articles_df is None:
        st.info("⚠️ Choose 'Articles' or 'Both' to fetch news content.")
    elif articles_df.empty:
        st.warning("No article results.")
    else:
        st.success(f"Found {len(articles_df)} articles")

        # Pagination
        page_size = 10
        total_pages = max(1, (len(articles_df) - 1) // page_size + 1)
        page = st.number_input("Page", 1, total_pages, 1)
        start = (page - 1) * page_size
        end = start + page_size

        for _, row in articles_df.iloc[start:end].iterrows():
            with st.container(border=True):
                st.markdown(
                    f"### {row['Title']}  \n"
                    f"<small>{row['Source']} — {row['Published']} &nbsp;&nbsp;|&nbsp;&nbsp; "
                    f"**Lang:** {article_lang_label}</small>",
                    unsafe_allow_html=True
                )
                if row.get("Summary"):
                    with st.expander("Summary", expanded=False):
                        st.write(row["Summary"])
                st.link_button("Open Article ↗", row["URL"])

        st.markdown("---")

        from io import BytesIO
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "⬇️ CSV",
                articles_df.to_csv(index=False).encode("utf-8"),
                "articles.csv",
                "text/csv"
            )

        with col2:
            st.download_button(
                "⬇️ TXT",
                export_articles_txt(articles_df).encode("utf-8"),
                "articles.txt",
                "text/plain"
            )

        with col3:
            doc = export_articles_docx(articles_df)
            if doc:
                buf = BytesIO()
                doc.save(buf)
                st.download_button(
                    "⬇️ Word",
                    buf.getvalue(),
                    "articles.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
